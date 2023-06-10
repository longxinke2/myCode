import re
import os
import wx
import sqlalchemy
import math
import functools
import numpy
import pandas as pd
from IPython.display import clear_output
from decimal import Decimal
import keyboard
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  #设置对象居中、对齐等。
from docx.enum.text import WD_TAB_ALIGNMENT,WD_TAB_LEADER,WD_LINE_SPACING  #设置制表符等
from docx.shared import Inches   #设置图像大小
from docx.shared import Pt,Cm   #设置像素、缩进等
from docx.shared import RGBColor    #设置字体颜色
from docx.shared import Length    #设置宽度
from docx.oxml.ns import qn  #设置中文版式
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32com.client
import pyautogui
from pyecharts.charts import Bar,Line,PictorialBar,Map,Pie,Sankey,Grid,Gauge,Radar,WordCloud,Liquid
from pyecharts import options as opts
from pyecharts.render import make_snapshot
from snapshot_selenium import snapshot
from pyecharts.commons.utils import JsCode
from pyecharts.components import Image
from pyecharts.options import ComponentTitleOpts
import seaborn as sns
import matplotlib.pyplot as plt
from IPython.display import display
from pyautogui import alert as printw

def bxwl2(x_group,y_group,max_n):#y_group是生源人数，落实人数，落实率
    p1 = (
        PictorialBar()                       #象形柱状图
        .add_xaxis(x_group)  #x轴
        .add_yaxis(
            "生源人数",              #系列名字
            y_group[0], #数据
            label_opts=opts.LabelOpts(is_show=True,
                                      formatter = "{c|{c}人}",rich={"c": {"fontSize": 21, "fontFamily": 'Times New Roman',"padding":[0,0,0,-60],},}), #数值标签
            symbol_size=40,              #每个图标的大小，重复图表则设置，不重复（类似山峰图）则去掉     
            symbol_repeat= True,        #是否重复图标
            is_symbol_clip=True,        #是否裁剪图标
            symbol=bd,                  #图标path
            symbol_offset=[-30, 0],
            color=f'{color_2}',
            yaxis_index=0)
        .add_yaxis(
            "落实人数",              #系列名字
            y_group[1], #数据
            label_opts=opts.LabelOpts(is_show=True,
                                      formatter = "{c|{c}人}",rich={"c": {"fontSize": 21, "fontFamily": 'Times New Roman',"padding":[0,0,0,70],},}), #数值标签
            symbol_size=40,              #每个图标的大小，重复图表则设置，不重复（类似山峰图）则去掉     
            symbol_repeat= True,        #是否重复图标
            is_symbol_clip=True,        #是否裁剪图标
            symbol=bd,                  #图标path
            symbol_offset=[30, 0],
            color=f'{color_1}',
            yaxis_index=0)
        .set_global_opts(yaxis_opts=opts.AxisOpts(max_=max_n*2,is_show = False), #去掉Y轴,设置最大轴值
                         xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(font_size = 21,font_family = '宋体',color = 'black')))#设置坐标轴文字样式               
        .set_global_opts(legend_opts=opts.LegendOpts(item_width=20,item_height=15,
                                                     textstyle_opts=opts.TextStyleOpts(font_size = 18,font_family = '宋体')))#调整图例格式
        .extend_axis(                  #增加扩展轴     
            yaxis=opts.AxisOpts(
                name="毕业去向落实率",
                type_="value",
                min_=0,
                max_=100,             #设置次坐标轴最大值
                position="right",
                is_show = False
                )))
    line = (                         #折线图
        Line()
        .add_xaxis(x_group) #x轴
        .add_yaxis(                  
            "毕业去向落实率",
            y_group[2],#数据
            label_opts=opts.LabelOpts(is_show=True,position='top',background_color=color_1,color='white',
                                      formatter=JsCode("function (params) {return params.value[1] + '%'}"),
                                      font_size = 26,font_family = 'Times New Roman'),
            yaxis_index=1,
            linestyle_opts=opts.LineStyleOpts(width=0),
            itemstyle_opts=opts.ItemStyleOpts(color=color_1),symbol='circle',symbol_size=6
        ))
    return p1.overlap(line)

def xy(datat,x,x1,y):
    return datat[datat[x]==x1][y].values[0]

def get_og_name(datat,is_p=False):#og=othergroupby
    return list(filter(lambda x:(is_p and '_proportion' or '_num') in x,(list(datat.columns))))

def three_word(start,loop_word_group,end):
    string1=start
    for j in loop_word_group:
        string1+=j
    return string1[:-1]+end

def sum_list(data,ziduan,lst,sum_ziduan):#is_p表示是否对占比求和
    if 'proportion' in sum_ziduan:
        return change_to_decimal(sum([change_to_decimal(x,rev=True) for x in data[data[ziduan].isin(lst)][sum_ziduan][data[sum_ziduan]!='-'].tolist()]))
    return data[data[ziduan].isin(lst)][sum_ziduan][data[sum_ziduan]!='-'].sum()

def bxwl1(x_group,y_group,per):  
    color_fuc = """
                    function (params) {
                        if (params.dataIndex == 0) return """ + f"'{color_2}';" + """if (params.dataIndex == 1)return """ + f"'{color_3}';" +"""
                        if (params.dataIndex == 2) return """ + f"'{color_4}';" +""" if (params.dataIndex == 3)return """ + f"'{color_5}';" +"""}   
                   """
    #柱形图        
    a = (
        Bar()
        .add_xaxis(x_group)
        .add_yaxis("",y_group,bar_width=35,stack="stack1",
                    label_opts=opts.LabelOpts(position='top',color=JsCode(color_fuc),formatter='{c}%',font_size = 23,font_family = 'Bahnschrift SemiLight Condensed'),
                    itemstyle_opts=opts.ItemStyleOpts(color=JsCode(color_fuc)))#设置柱形图样式
        .set_global_opts(yaxis_opts=opts.AxisOpts(is_show=False),#不显示y轴
                        xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(font_size = 16,font_family = '宋体'))))


    #水波图
    b = (
        Liquid()
        .add("毕业去向落实率", [change_to_decimal(per,True)/100], is_outline_show=False,shape = 'circle',center=["25%", "50%"],is_animation=False,
             color = [get_lighter_color('#FFFFFF',color_1)[0][2]],background_color=['#FFFFFF'],
             label_opts=opts.LabelOpts(font_size=25,position="inside",color = color_1,formatter="{c|%s}" % per +'\n'*10+'{d|毕业去向落实率}',
                                       rich={"c": {"fontSize": 27, "fontFamily": 'Bahnschrift SemiLight Condensed',"fontWeight":'bolder'},
                                             "d": {"fontSize": 18, "fontFamily": '宋体'}}))
        .set_series_opts(radius="60%"))

    c = (
        Liquid()
        .add("", [1], is_outline_show=False,shape = bd,center=["25%", "50%"],is_animation=False,color = [color_1],
            label_opts=opts.LabelOpts(is_show=False))
        .set_series_opts(radius="40%"))

    #组合在一起
    grid = (
         Grid(init_opts = opts.InitOpts(height = '350px',width='1000px'))
        .add(a, grid_opts=opts.GridOpts(pos_left="45%"))
        .add(b, grid_opts=opts.GridOpts(pos_left=""))
        .add(c, grid_opts=opts.GridOpts(pos_left=""))
    )
    return grid

def reduce_by_1(string):
    return change_to_decimal(100-float(re.sub('%','',string)))+'%'

# 取data的前三，可自动忽略汇总行和总计行，忽略非数字
def get_3rd(data,num,not_3=3):#num为倒序排序的依据,不取前三也行
    datat=data.copy()
    datat= datat[(datat.iloc[:,0] != '汇总') & (datat.iloc[:,0] != '总计') & (datat.iloc[:,1] != '汇总') & (datat.iloc[:,1] != '总计')]
    datat['temp']=data.apply(lambda x:0 if isinstance(x[num],str) else x[num],axis=1)
    datat=datat.sort_values(by='temp',ascending=False).reset_index(drop=True).drop(['temp'],axis=1)
    return [datat.loc[x].tolist() for x in range(not_3)]

def my_liquid(bys_count,group):
    p1_1 = Liquid()
    group1=sum100([x[1] for x in group])
    for index,i in enumerate(group):
        if len(group)>1:
            formatter="%d"%(i[1])+'人'+'\n'+group1[index]+'\n'*9+i[0]
        else:
            formatter="%d"%(i[1])+'人'+'\n'+"%.2f"%(i[1]/bys_count*100)+'%'+'\n'*9+i[0]
        p1_1.add("", [i[1]/bys_count], is_outline_show=False,shape = eval(L(i[0])),center=i[2],is_animation=False,color = [eval(f'color_{index+1}')],background_color='#D3D3D3',
             label_opts=opts.LabelOpts(font_size=25,position="inside",color = eval(f'color_{index+1}'),
            formatter=f'{formatter}'))
    p1_1.set_series_opts(radius="35%")
    return p1_1

def my_pie(group,title='',js='',radius=["30%", "55%"],center=["25%", "55%"]):
    group2=[x[1] for x in group]
    group1=sum100(group2)
    group=[*zip([x[0] for x in group],[*zip(group2,group1)])]
    b = (
        Pie()
        .add(
            "",group,radius =radius , #radius为饼图的半径，数组的第一项是内半径，第二项是外半
               center =center )               #设置饼状图位置，第一个百分数调水平位置，第二个百分数调垂直位置
        .set_global_opts(
            title_opts=opts.TitleOpts(title = title,pos_left='18%',pos_top='45%',
                                      title_textstyle_opts=opts.TextStyleOpts(font_size = 21,font_family = '宋体',color=color_1)), #标题字体样式配置项
            legend_opts=opts.LegendOpts(is_show=False),
        )
        .set_series_opts(label_opts=opts.LabelOpts(is_show=False))
    )
    if js:
        b = (
             Pie()
            .add("",group,radius=radius,center=center,
                    label_line_opts=opts.PieLabelLineOpts(length_2=40,linestyle_opts=opts.LineStyleOpts(width=2.5)),
                    label_opts=opts.LabelOpts(interval=0,
                        position="outside",
                        formatter=JsCode(
                        f"function(params) {js}"),
                        rich={
                            "b": {"fontSize": 18, "fontFamily": '宋体',"fontWeight":'normal',"padding":[0,0,-53,0]},
                            "c": {"fontSize": 18, "fontFamily": 'Bahnschrift SemiLight Condensed',"fontWeight":'normal'},
                            "d": {"fontSize": 30, "lineHeight": 10, "fontFamily": 'Bahnschrift SemiLight Condensed',"fontWeight":'normal'},
                        })
                 ) 
            .set_global_opts(
                title_opts=opts.TitleOpts(title = title,pos_left='18%',pos_top='45%',
                                          title_textstyle_opts=opts.TextStyleOpts(font_size = 21,font_family = '宋体',color=color_1)), #标题字体样式配置项
                legend_opts=opts.LegendOpts(is_show=False))#去掉图例
        )
    return b

#竖的，多个不同的列
def my_bar(x_group,y_group):      
    a = Bar().add_xaxis(x_group)
    for index,i in enumerate(y_group):
        a.add_yaxis(i[0],i[1],bar_width=35,label_opts=opts.LabelOpts(position='top',color=eval(f'color_{index+1}'),formatter='{c}',font_size = 23,font_family = 'Bahnschrift SemiLight Condensed'),itemstyle_opts=opts.ItemStyleOpts(color=eval(f'color_{index+1}')))#设置柱形图样式
    a.set_global_opts(yaxis_opts=opts.AxisOpts(is_show=False),xaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(font_size = 16,font_family = '宋体')),legend_opts=opts.LegendOpts(pos_right="20%")).set_colors([f'{color_1}',f'{color_2}',f'{color_3}',f'{color_4}',f'{color_5}'])
    return a

#横的，一个行
def my_bar2(x_group,y_group):
    jscode = "function (params) {return params.name" + " + '  ' + params.value + '人'}"
    b = (
        Bar()
        .add_xaxis(x_group)
        .add_yaxis("",y_group,bar_width=20)
        .reversal_axis()
        .set_series_opts(label_opts=opts.LabelOpts(position="right",color=color_2,font_weight='bold',
                                                   formatter= JsCode(jscode),font_size = 16,font_family = 'Bahnschrift SemiLight Condensed'),
                         itemstyle_opts=opts.ItemStyleOpts(color=f'{color_2}'))
        .set_global_opts(xaxis_opts=opts.AxisOpts(is_show=False),
                         yaxis_opts=opts.AxisOpts(axislabel_opts=opts.LabelOpts(font_size = 18,font_family = '宋体',color='black'),
                                                 axisline_opts=opts.AxisLineOpts(linestyle_opts=opts.LineStyleOpts(color=color_2)),
                                                 axistick_opts=opts.AxisTickOpts(linestyle_opts=opts.LineStyleOpts(color=color_2)))) #设置坐标轴文字样式
        .set_colors([f'{color_1}',f'{color_2}']))
    return b

def merge_pic(a,b,height='450px',width='1000px'):
    #组合在一起
    grid = (
         Grid(init_opts = opts.InitOpts(height =height ,width=width))
        .add(a, grid_opts=opts.GridOpts(pos_left="50%"))
        .add(b, grid_opts=opts.GridOpts(pos_left="")))
    return grid

def add_pic(html_path):
    make_snapshot(snapshot,html_path, 'result/temp.png',pixel_ratio = 5)
    a=document.add_paragraph()
    b=a.add_run().add_picture('result/temp.png',width=Cm(15.24))
    a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    os.remove(html_path)
    os.remove('result/temp.png')

def my_china_map(data_zip,label_fontsize=13):
    color = get_lighter_color('#FFFFFF',color_1)
    a= (
    Map()
    .add("",data_zip,"china",is_map_symbol_show = False,is_roam = False,#数据传入,地图类型设置:中国地图"china"
         name_map=dict_sf,
         label_opts=opts.LabelOpts(formatter= '{b}',font_size=label_fontsize,font_family = '宋体'),#标签设置
         itemstyle_opts=opts.ItemStyleOpts(border_color = '#A6A6A6')) #图形的描边颜色
    .set_global_opts(
         visualmap_opts=opts.VisualMapOpts(is_show=False,max_ = len(data_zip),range_color = color))#颜色渐变设置
    )
    a = (
         Grid(init_opts = opts.InitOpts(height = '700px',width='1000px'))  #设置图形大小,高度:height,宽度:width
        .add(a, grid_opts=opts.GridOpts(pos_left='35%'))#位置调整
    )
    path = a.render()
    with open(path, 'r') as file:
        html = file.read()
    with open(path, 'w') as file:
        html=re.sub('https://assets.pyecharts.org/assets/maps/china.js','model/china.js',html)
        file.write(html)
    display(a.render_notebook())
    return path

def find_province(y,is_all=True):#is_all 是找全称还是简称
    zidian = {'西藏' : '西藏自治区',
 '内蒙古' : '内蒙古自治区',
 '广西' : '广西壮族自治区' ,
 '新疆' :'新疆维吾尔自治区' ,
 '宁夏' :'宁夏回族自治区',
 '香港' :'香港特别行政区',
 '澳门' :'澳门特别行政区',
 '北京' :'北京市',
 '上海' :'上海市' ,
 '天津' :'天津市' ,
 '重庆' :'重庆市',
 '河北' :'河北省',
 '山东' :'山东省',
 '辽宁' :'辽宁省',
 '黑龙江' :'黑龙江省',
 '甘肃' :'甘肃省',
 '吉林' :'吉林省',
 '青海' :'青海省',
 '河南' :'河南省',
 '江苏' :'江苏省',
 '湖北' :'湖北省',
 '湖南' :'湖南省',
 '浙江' :'浙江省',
 '江西' :'江西省',
 '广东' :'广东省',
 '云南' :'云南省',
 '福建' :'福建省',
 '台湾' :'台湾省',
 '海南' :'海南省',
 '山西' :'山西省',
 '四川' :'四川省',
 '陕西' :'陕西省',
 '贵州' :'贵州省',
 '安徽' :'安徽省'}
    return is_all and zidian[list(filter(lambda x:x in y,zidian))[0]] or list(filter(lambda x:x in y,zidian))[0]

def write_zw(content,is_show=False):#传入一个字符串格式，将你要写的题注或者图注放入即可
    document.add_paragraph(content,style='正文段落文本')
    if is_show:print(content)
    
def write_tz(content,is_show=False):#传入一个字符串格式，将你要写的题注或者图注放入即可
    document.add_paragraph(content,style='标准题注')
    if is_show:print(content)

def open_word(path):
    # 主要用于查看改动的效果
    # 创建Word应用程序对象
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = True
    path=os.path.abspath(path)# 不知道为什么当前路径不行，要完整路径
    doc = word.Documents.Open(path)
    word.WindowState = 3
    word.Activate() # 置于所有应用之前

def add_head(num,content):
    document.add_heading('',level = num).add_run(content)

def sd(look=True):#sd=save doc
    dc_path='result/test.docx'
    global document
    document.save(dc_path)
    if look:open_word(dc_path)
    
def read_doc(moban):
    path=f'model/{moban}.docx'
    
    global color_1
    global color_2
    global color_3
    global color_4
    global color_5
    if moban=='红色模板':
        color_1 = '#147BC5' #主配色
        color_2 = '#FBB03E' #主配色
        color_3 = '#24AAE1' #主配色
        color_4 = '#E3798D' 
        color_5 = '#00ABCB'
    try:
        global document
        document=Document(path)
    except:
        raise Exception('路径似乎出错了')

def get_3_data(xxmc,year,string,type1=''):#type1是省份库标识
    conn = get_yjy_db('dw_database')
    sql = ''
    type2 = type1 and f'_{type1}_province' or type1
    if string=='派遣':
        sql =f"SELECT * FROM dw_s_employment_auto where xxmc = '{xxmc}' and substring(bynd,1,4) = '{year}'"
    if string=='单位调研':
        sql =f"SELECT * FROM dw_{year}_c_short{type2}_survey_auto where dy_xxmc = '{xxmc}' "
    if string=='学生调研':
        sql =f"SELECT * FROM dw_{year}_s_short{type2}_survey_auto where dy_xxmc = '{xxmc}' "
    if sql=='':
        raise Exception('请输入正确的数据类型')
    data = pd.read_sql(sql=sql, con=conn)
    return data

def get_normal_data(sql,db):
    conn = get_yjy_db(db)
    data = pd.read_sql(sql=sql, con=conn)
    return data

def add_table(data1,need_1st_merge=False):#need_1st_merge:第一列需要合并处理
    global document
    data=data1.copy()
    if need_1st_merge:
        cn=data.columns[0]#cn=columns.name
        cn1=data.columns[1]
        mask=(data[cn]!='总计')&((data[cn]==data[cn].shift(1))|(data[cn].shift(1)==L(cn))|(data[cn1].shift(1)=='汇总'))&(data[cn1]!='汇总')
        lst = mask.tolist()
        result=[]
        temp=[]
        for i in range(len(lst)):
            if lst[i]:temp.append(i)
            if not lst[i] and temp:
                result.append(temp)
                temp=[]
        group1=[]
        for i in get_og_name(data,True)+['proportion']:
            for j in result:
                group1.append(change_to_decimal(sum(data.loc[j][i].apply(lambda x:change_to_decimal(x,True)).tolist()))+'%')
        data.loc[data1[cn1]=='汇总',i]=group1

    table = document.add_table(rows=data.shape[0],cols=data.shape[1],style="表格-全部")#建立表格
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            if str(data.iloc[i, j])=='left':
                table.cell(i,j).merge(table.cell(i,j-1))
                continue
            if str(data.iloc[i, j])=='up':
                table.cell(i,j).merge(table.cell(i-1,j))
                continue
            if str(data.iloc[i, j])=='汇总':
                table.cell(i,j).merge(table.cell(i,j-1))
                table.cell(i,j-1).text = str(data.iloc[i, j-1])+' 汇总'
                continue
            table.cell(i,j).text = str(data.iloc[i, j])
    #处理第一列
    if need_1st_merge:
        for i in result:
            if len(i)>1:
                table.cell(i[0],0).merge(table.cell(i[-1],0))
                table.cell(i[0],0).text= str(data.iloc[i[0],0])
    # table.alignment = WD_TABLE_ALIGNMENT.CENTER   #设置整个表格居中
#     table.cell(0,0).merge(table.cell(1,0))
    #调整单元格内部的格式
    for a in table.rows:
        a.height = Cm(0.5)
        for b in a.cells:
            b.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            b.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for c in b.paragraphs:
                c.paragraph_format.first_line_indent = Pt(0)
                c.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
                c.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
                c.paragraph_format.line_spacing=1
                c.paragraph_format.space_before=Pt(0)
                c.paragraph_format.space_after=Pt(0)
                for item in c.runs:
                    item.font.name = 'Times New Roman'  # 英文字体设置
                    item._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    item.font.size=Pt(9)
                    item.font.no_proof=True
                    item.font.no_proof=True#无视系统检索报错的下划线
    document.add_paragraph().add_run('')#增加空行
    
def admin():
    while True:
        keyboard.wait('F2')
        global baba
        baba = not baba
        printw(baba and '开发者模式已开启，我将以高达形态出击！' or '已关闭开发者模式')

def L(string):
    try:
        return index_dict[string]
    except:
        keys = {v: k for k, v in index_dict.items()}  # 创建反向查找的字典
        key = keys.get(string)  # 查找该值对应的键
        if key is None:
            raise Exception('没有这个字段')
        return key

# 对data按照string分组+对行排序
def get_group_order_table(data,string,string1='',other_groupby='',order=True,need_sum=True,need_title=True):
    # 获取分组数据
    other_groupby = other_groupby or (baba and L(input('输入额外的分组维度') or '学历') or '')
    data1 = groupby_data(data,string,string1,other_groupby)
    if string1!='':
        data4=groupby_data(data,string,other_groupby=other_groupby)
        data4[string1]='汇总'
        data1 = pd.concat([data1,data4], ignore_index=True)
    group = {'default':'默认排序','super':f'按照给定的{L(string)}顺序排序','num_down':'按照人数降序','num_up':'按照人数升序'}
    if baba:
        # 排序
        if len(xy_order_str)==0:
            result = get_order_way('排序规则：',group)
        else:
            result = f'按照给定的{L(string)}顺序排序'
        data1 = order_table(data1,result,string,group,string1)
        if data1 is None:
            return 
    else:
        #默认按照人数降序
        result = string=='jylbfl' and '按照给定的就业类别分类顺序排序' or '按照人数降序'
        data1 = order_table(data1,result,string,group,string1)
        
    if need_title:
        if other_groupby=='xl':
            group = {'本':'本科毕业生','硕':'毕业研究生','研':'毕业研究生','博':'毕业研究生','专':'专科毕业生'}
            f = lambda y: list(filter(lambda x:x in y,list(group.keys()))) and group[list(filter(lambda x:x in y,list(group.keys())))[0]] or '全体毕业生'
            if string1=='':
                dict2={x:[L(string),'up'] if x==string else 'num' in x and [f(x),'人数'] or ['left','占比'] for x in data1.columns}
            else:
                dict2={x:[L(x)] if x==string or x==string1 else 'num' in x and [f(x)] or ['left'] for x in data1.columns}
            data2 = pd.DataFrame(dict2)
        else:
            data2 = pd.DataFrame({x:[L(string)] if x==string or x==string1 else 'num' in x and ['人数'] or ['占比'] for x in data1.columns})
        
    if need_sum:
        dict3={x:'总计' if x==string else(x==string1 and 'left' or 'num' in x and [data1[x].sum()] or ['100.00%']) for x in data1.columns}
        data3 = pd.DataFrame(dict3)
        
    #合并
    if need_title:data1 = pd.concat([data2,data1], ignore_index=True)
    if need_sum:data1 = pd.concat([data1,data3], ignore_index=True)
        
    data1.replace(0,'-',inplace=True)
#     clear_output()  # 要不要打印出来自己在外面决定吧
#     display(data1)
    return data1

####################################以下为中间函数，最好不使用######################################################################

def RGB_to_Hex(rgb):
    RGB = rgb.split(',') # 将RGB格式划分开来
    color = '#'
    for i in RGB:
        num = int(i) # 将R、G、B分别转化为16进制拼接转换并大写 hex() 函数用于将10进制整数转换成16进制，以字符串形式表示
        color += str(hex(num))[-2:].replace('x', '0').upper()
    return color  # RGB格式颜色转换为16进制颜色格式

def RGB_list_to_Hex(RGB):
    # RGB = rgb.split(',') # 将RGB格式划分开来
    color = '#'
    for i in RGB:
        num = int(i)
        # 将R、G、B分别转化为16进制拼接转换并大写 hex() 函数用于将10进制整数转换成16进制，以字符串形式表示
        color += str(hex(num))[-2:].replace('x', '0').upper()
    return color
def Hex_to_RGB(hex):
    r = int(hex[1:3], 16)
    g = int(hex[3:5], 16)
    b = int(hex[5:7], 16)
    rgb = str(r) + ',' + str(g) + ',' + str(b)
    return rgb, [r, g, b]
def gradient_color(color_list, color_sum=700):
    color_center_count = len(color_list)
    color_sub_count = int(color_sum / (color_center_count - 1))
    color_index_start = 0
    color_map = []
    for color_index_end in range(1, color_center_count):
        color_rgb_start = Hex_to_RGB(color_list[color_index_start])[1]
        color_rgb_end = Hex_to_RGB(color_list[color_index_end])[1]
        r_step = (color_rgb_end[0] - color_rgb_start[0]) / color_sub_count
        g_step = (color_rgb_end[1] - color_rgb_start[1]) / color_sub_count
        b_step = (color_rgb_end[2] - color_rgb_start[2]) / color_sub_count
        # 生成中间渐变色
        now_color = color_rgb_start
        color_map.append(RGB_list_to_Hex(now_color))
    for color_index in range(1, color_sub_count):
        now_color = [now_color[0] + r_step, now_color[1] + g_step, now_color[2] + b_step]
        color_map.append(RGB_list_to_Hex(now_color))
        color_index_start = color_index_end
    return color_map

def get_lighter_color(*args):
    colors = gradient_color(args)
    return colors[::int(len(colors)/10)][:]

def change_to_decimal(x,rev=False):#rev=由string转回float
    return rev and float(re.sub('%','',x)) or str(Decimal(x).quantize(Decimal("0.00")))

# 对data按照string分组
def groupby_data(data,string,string1='',other_groupby=''): 
    groupt= [string,string1] if len(string1)!=0 else [string]
    data1 = data.groupby(groupt)['xxdm'].count().reset_index(name='num')  # 计算组别数量
    data1['proportion'] = sum100(data1['num'].tolist())   # 计算占比
    if other_groupby!='':
        data2 = data.groupby(groupt+[other_groupby])['xxdm'].count().reset_index(name='num')  
        group = data2[other_groupby].unique()
        if other_groupby=='xl':
            group = sorted(group,key=lambda x: ('本' in x, '研' in x, '专' in x),reverse=True)
        for i in group:
            data3 = data2[data2[other_groupby]==i]
            data3['proportion'] = sum100(data3['num'].tolist())  # 计算占比
            data3.columns = groupt+[i,f'{i}_num',f'{i}_proportion']
            data1 = data1.merge(data3,how='left',on=groupt)
        # 转化为int
        for i in data1.columns:
            if 'num' in i:
                data1[i] = data1[i].apply(lambda x: not math.isnan(x) and int(x) or 0)
        data1.drop(data2[other_groupby].unique(),axis=1,inplace=True)
    data1.replace(numpy.nan,0,inplace=True)#nan值替换成0
    data1 = data1.reindex(columns=data1.columns.tolist()[:len(groupt)]+data1.columns.tolist()[len(groupt)+2:]+data1.columns.tolist()[len(groupt):len(groupt)+2])
    return data1

# 连接研究院数据库
def get_yjy_db(db):
    return sqlalchemy.create_engine(f"mysql+pymysql://yjy_user:Yjy123456@am-wz9el267w54i2r7ip131930o.ads.aliyuncs.com:3306/{db}")

# 对table的行按照排序选择的result排序
def order_table(table,result,string,group={},string1=''):
    a = list(filter(lambda x:group[x]==result,group.keys()))
    if len(a)==0 or a[0]=='default':
        return table
    table['汇总']=table[string1!='' and string1 or string].apply(lambda x:x=='汇总' and 2 or 1)#string1为空代表不需要汇总，随便搞了
    if a[0]=='super' :
        if string=='jylbfl':
            order_str=jylbfl_order_str
        else:
            global xy_order_str
            xy_order_str = xy_order_str if xy_order_str!='' else input('请输入排序规则：')
            order_str=xy_order_str
        list1 = []
        try:
            for i in list(table[string]):
                list1.append(re.search(i, order_str).span(0)[0])
        except:
            print(f'未在表格中找到{L(string)}字段')
            if string=='xl':xy_order_str=''
            return
        table['order'] = list1
        table.sort_values(by=string1!='' and ['order','汇总'] or 'order',inplace=True)
        table.drop('order',axis=1,inplace=True)
    if a[0]=='num_down':
        table['汇总']=2-table['汇总']
        table.sort_values(by=string1!='' and [string,'汇总','num'] or 'num',inplace=True,ascending=False)
    if a[0]=='num_up':
        table.sort_values(by=string1!='' and [string,'汇总','num'] or 'num',inplace=True)
    table.drop('汇总',axis=1,inplace=True)
    return table

# 完美百分数
def sum100(arr1):
    arr=arr1.copy()
    # 使用哈希表记录每个元素出现的次数
    freq = {}
    for num in arr:
        if num in freq:
            freq[num] += 1
        else:
            freq[num] = 1
    # 创建一个结果列表
    candidates = []
    # 遍历哈希表，将每个元素和它出现的次数添加到结果列表中
    
    for key, value in freq.items():
        candidates.append([key, int(round(key/sum(arr),5)*10000), key/sum(arr) *10000 - int(round(key/sum(arr),5)*10000), value])
    target = 10000
    for i in candidates:
        target-=i[1]*i[3]
    if target == 0:
        return [change_to_decimal(x/sum(arr)*100)+'%' for x in arr]
    result = []

    def dfs(combination, current_sum, start):
        # 如果当前组合的数字总和超过了目标值，则停止递归
        if current_sum > target:
            return
        # 如果当前组合的数字总和等于目标值，则将其添加到结果中
        if current_sum == target:
            result.append(combination)
            return
        # 递归搜索剩余的数字
        for i in range(start, len(candidates)):
            if len(result)==1 and flag:break
            dfs(combination + [candidates[i]], current_sum + candidates[i][3], i + 1)
    
    # 超过一定长度的数组拿到一个结果就撤，不然时间复杂度太高了
    flag =False
    if len(arr)>10:
        flag=True

    dfs([], 0, 0)
    temp_g=[]
    for r in result:
        sum_num = 0
        for r1 in r:
            sum_num+=r1[2]*r1[3]
        temp_g.append(sum_num)
    tgroup = []
    for i in range(0,len(temp_g)):
        if temp_g[i]==max(temp_g):
            tgroup=result[i]
            break
    if tgroup==[]:
        print('该处百分比相加不为1')
    for j in tgroup:
        candidates[candidates.index(j)][1] += 1
    for i in range(len(arr)):
        for j in candidates:
            if arr[i]==j[0]:
                arr[i]=j[1]
    return [change_to_decimal(x/100)+'%' for x in arr]

class btnFrame(wx.Frame):
    def __init__(self, parent,title, btn_group,callback):
        wx.Frame.__init__(self, parent, title='', size=(10*max(len(v) for v in btn_group.values())+80, len(btn_group)*30+100))
        
        # 回调函数
        self.callback=callback
        
        sizer = wx.BoxSizer(wx.VERTICAL)
        self.m_staticText1 = wx.StaticText( self, wx.ID_ANY, f"{title}", wx.DefaultPosition, wx.DefaultSize, 0 )
        self.m_staticText1.Wrap( -1 )

        sizer.Add( self.m_staticText1, 0, wx.ALL, 5 )

        for i in btn_group:
            exec(f"self.{i}=wx.Button(self, label='{btn_group[i]}')") # 创建btn
            exec(f"sizer.Add(self.{i}, 0, wx.ALL, 5)") # 添加进
            exec(f"self.{i}.Bind(wx.EVT_BUTTON, lambda event,self=self, i=btn_group[i]: (self.callback(i),self.Destroy()))")

        self.SetSizer(sizer)
        
        # 初始化窗口
        self.Centre()
        self.Show()
        self.SetWindowStyle(wx.DEFAULT_FRAME_STYLE | wx.STAY_ON_TOP)
        
# 弹出frame获取排序选择
def get_order_way(title,btn_group):
    try:
        del app
    except:
        pass
    selection_value = ''
    def getvalue(value):
        nonlocal selection_value
        selection_value = value
    app = wx.App()
    frame = btnFrame(None,title,btn_group,getvalue)
    app.MainLoop()
    return selection_value

####################################################### 全局变量##############################################################
baba = False
document = ''
read_doc('红色模板')
xy_order_str = ''#学院顺序缓存
jylbfl_order_str = '协议和合同就业继续深造灵活就业自主创业待就业暂不就业'

color_1 = '#147BC5' #主配色
color_2 = '#FBB03E' #主配色
color_3 = '#24AAE1' #主配色
color_4 = '#E3798D' 
color_5 = '#00ABCB'

dict_sf={'河北省': '河北', '香港特别行政区': '香港', '天津市': '天津', '山西省': '山西', '内蒙古自治区': '内蒙古', '四川省': '四川', '海南省': '海南', '江西省': '江西', '北京市': '北京', '黑龙江省': '黑龙江', '云南省': '云南', '西藏自治区': '西藏', '浙江省': '浙江', '广西壮族自治区': '广西', '辽宁省': '辽宁', '澳门特别行政区': '澳门', '福建省': '福建', '安徽省': '安徽', '青海省': '青海', '贵州省': '贵州', '吉林省': '吉林', '湖南省': '湖南', '陕西省': '陕西', '台湾省': '台湾', '山东省': '山东', '江苏省': '江苏', '重庆市': '重庆', '广东省': '广东', '湖北省': '湖北', '甘肃省': '甘肃', '新疆维吾尔自治区': '新疆', '宁夏回族自治区': '宁夏', '上海市': '上海', '河南省': '河南'}

# 字典
index_dict = {# 其他：
            '湖南':'hunan','贵州':'guizhou','boy':'男生','girl':'女生',
            #基础指标
            'myd':'满意','ppd':'匹配','xgd':'相关','tjd':'愿意','gzd':'关注','xb':'性别','xl':'学历','yx':'学院','zy':'专业','sysf':'生源省份',
            'mz':'民族','jyq':'乐观','zqj':'乐观','zsp':'满意','jjx':'满意','lzt':'乐观','sjy':'满意','yjy':'满意','ljy':'满意','jylbfl':'就业类别分类','lsbyqx':'毕业去向落实','jylb':'就业类别','jylbfl':'就业类别分类',
            #派遣数据
            'sxgxlx':'升学院校类型','sxgxbq':'升学院校层次',
            #分数指标
            '5':'非常高','4':'高','3':'比较高','2':'不太高','1':'不高',
            #教育教学调研
            'dy_cyzyxgd':'创业专业相关度','dy_cyyy':'创业原因','dy_cylx':'创业类型','dy_cyfs':'创业方式',
              'dy_cykn':'创业困难','dy_cyfwjy':'创业服务建议','dy_mxmyd':'母校满意度','dy_mxtjd':'母校推荐度','dy_jyjx':'教育教学',
              'dy_mxtjd':'母校推荐度','dy_jyjx_tsjy':'通识教育满意度','dy_tsjyjy':'通识教育建议','dy_jyjx_zyjy':'专业教育满意度','dy_zykcjy':'专业教育建议',
             'dy_jyjx_sjjx':'实践教学满意度','dy_sjjxjy':'实践教学建议','dy_zynljy':'职业能力教育建议','dy_zzdgjjy':'师资水平建议',
             'dy_jyjx_zynljy':'职业教育满意度','dy_jyjx_szsp':'师资水平','dy_jyjx_xsjz':'学术讲座','dy_jyjx_jxss':'教学设施',
            'dy_jyjx_khfs':'考核方式','dy_sjjxmyd':'实践教学满意度','dy_jyjx_zhmyd':'综合满意度',
            'dy_sjjx_kcsy':'课程实验','dy_sjjx_kcsyx':'课程实用性','dy_sjjx_sjkbz':'实践课比重','dy_sjjx_bylw':'毕业论文',
             'dy_sjjx_sys':'实验室的使用与管理','dy_kcszjy':'课程设置建议','dy_szsp':'师资水平满意度','dy_szsp_jxtd':'教学态度满意度',
             'dy_szsp_jxnr':'教学内容满意度','dy_szsp_jxff':'教学方法、方式满意度','dy_qzwt':'求职问题','dy_zxjl':'在校经历',
             'dy_szsp_sshd':'师生互动满意度','dy_szsp_zynl':'专业知识能力满意度',
             'dy_jycyfwmyd':'就业创业服务满意度','dy_jycyfwjy':'就业创业服务建议','dy_jycy':'就业创业',
             'dy_jycy_zcxc':'政策宣传与讲解','dy_jycy_xxtg':'信息提供与发布','dy_jycy_zyzx':'职业选择咨询/辅导','dy_jycy_jnpx':'就业/创业技能培训','dy_jycy_xlts':'求职心理调适','dy_jycy_mszd':'面试指导与训练',
             'dy_jycy_sxzd':'升学指导','dy_jycy_lxzd':'留学指导','dy_jycy_kc':'就业/创业课程','dy_jycy_xssxh':'线上双选会','dy_jycy_xyzp':'校园招聘活动','dy_jycy_jysx':'就业手续办理','dy_jycy_knbf':'就业困难群体帮扶',
              #就业调研
              '自主创业':'zzcy','dwhy':'单位行业','zzcy':'自主创业','sfmc':'所在省份','dy_lzcs':'离职次数','dy_jyqjyq':'专业前景预期',
              'dy_jymyd':'就业满意度','dy_jyzyxgd':'就业专业相关度','dy_jyjxmyd':'教育教学满意度','dy_jyqs':'就业歧视',
              'dy_jyzybxgyy':'就业专业不相关原因','dy_zwppd':'职业期待匹配度','dy_xzmyd':'薪资满意度','dy_jyyx':'月薪',
              'dy_jypjxz':'平均月薪','dy_shbz':'社会保障','dy_gwfzqj':'发展前景预期','dy_qwyz':'期望月薪',
              'dy_shbzmyd':'社会保障满意度','dy_dwpj':'岗位评价','dy_dwzmd':'单位知名度','dy_pxjh':'培训机会','dy_jskj':'晋升空间',
              'dy_gzhj':'工作环境','dy_gzzzx':'工作自主性','dy_gzyl':'工作压力','dy_gzwdx':'工作稳定性','dy_gwfzqj':'岗位发展前景',
              'dy_gzdd':'工作地点','dy_qywh':'企业文化','dy_gzfw':'工作氛围','dy_dwzhmyd':'单位综合满意度',
              'dy_lzyy':'离职原因','dy_gzys':'求职关注因素','dy_gzys_xcsp':'薪酬水平','dy_gzys_shbz':'社会保障',
              'dy_gzys_gzwdx':'工作稳定度','dy_gzys_gzcs':'工作城市','dy_gzys_dwsw':'单位社会声望','dy_gzys_gzhj':'工作环境',
              'dy_gzys_fzkj':'发展空间','dy_gzys_jrqw':'父母家人期望','dy_gzys_rjgx':'人际关系','dy_gzys_zydk':'专业对口',
              'dy_gzys_grxq':'个人兴趣','dy_gzys_rczc':'人才政策','dy_qztj':'求职途径','dy_offer':'offer数','dy_qzsc':'求职时长',
              'dy_wjyyy':'未就业原因','dy_qwdwlx':'期望单位类型','dy_qwjyss':'期望就业省市','dy_cgcjyy':'出国出境原因',
              'dy_qwbf':'期望帮扶','dy_nltsxq':'期望求职帮助','dy_jxszzyxgd':'继续深造与专业相关度','dy_gnsxyy':'国内升学原因',
              'dy_yqyx_zygh':'职业规划影响','dy_yqyx_xlzt':'就业乐观度',
              #单位调研
              'cdy_dwgm':'单位规模','cdy_dwxz':'单位性质','cdy_dwhy':'单位行业','cdy_dwszd':'单位所在地','cdy_zpqd':'招聘渠道',
              'cdy_kzys':'关注因素','cdy_zpxz':'招聘薪酬范围','cdy_zygzd':'专业关注度','cdy_rcpymyd':'人才培养认同度',
              'cdy_bysmyd':'毕业生满意度','cdy_byspj_sxpz':'思想品质','cdy_rcpyjy':'人才培养建议','cdy_lxzpyy':'来校招聘原因',
              'cdy_byspj_gztd':'工作态度','cdy_byspj_zysp':'专业水平','cdy_byspj_zynl':'职业能力','cdy_byspj_fzql':'发展潜力',
              'cdy_byspj':'毕业生评价','cdy_gwsysj':'岗位适应时间','cdy_lzqk':'离职情况','cdy_lzyy':'离职原因','cdy_zpkn':'招聘困难',
              'cdy_jyfwmyd':'就业服务满意度','cdy_zmyjyfw':'就业服务','cdy_jyfwjy':'就业服务建议'}

#矢量图设置,可以在阿里巴巴矢量图网站自行查找需要的图标,复制SVG代码
bd = 'path://M945.08544 366.6432l-425.97376-183.1936a20.5312 20.5312 0 0 0-8.12544-1.6896c-2.7648 0-5.52448 0.5632-8.14592 1.6896L76.87168 366.6432c-8.98048 3.84512-14.92992 13.65504-14.92992 24.59136 0 10.9568 5.94944 20.74112 14.92992 24.60672L502.8352 599.04a20.29056 20.29056 0 0 0 16.27136 0l331.1872-142.43328v145.7152c-17.75104 9.40032-30.08512 29.8496-30.08512 53.62688 0 23.7824 12.33408 44.22656 30.08512 53.62688v50.83136h43.6224v-50.82112c17.76128-9.40032 30.1056-29.8496 30.1056-53.63712s-12.34432-44.2368-30.1056-53.63712V437.84704l51.16416-22.00576c8.97536-3.8656 14.9248-13.65504 14.9248-24.60672 0-10.93632-5.94944-20.74624-14.9248-24.59136z m-434.0992 269.16352a47.96928 47.96928 0 0 1-19.03104-3.95776L226.4064 517.632v157.70112c0 129.16224 193.65888 167.936 254.208 167.936h60.69248c45.35808 0 254.22848-38.77376 254.22848-167.936v-157.696l-265.55904 114.21184a47.76448 47.76448 0 0 1-18.9952 3.95776z'
sf = 'path://M0,10 L10,10 C5.5,10 5.5,5 5,0 C4.5,5 4.5,10 0,10 z'
male = 'path://M724.8 314.2c2.8 1.6 5.5 3.2 8.2 5.1l8.2 5.6 0.1 9.8c0.4 28.5-3 52.6-10.4 71.1-6.9 17-17.1 29.7-30.7 37.3-9.3 33.1-20.2 63.8-35.6 89.6-17.8 29.7-41.2 52.5-74.4 64.8-14.8 5.4-52.7 8-89.1 7.4-35.2-0.5-70.5-4-84.5-10.5-30.1-13.9-51.3-37.2-67.4-66.3-13.8-24.8-23.6-53.7-32.1-84.2-14.1-7.4-24.7-20-31.8-37.1-7.8-18.7-11.4-43.2-10.9-72.2l0.1-9.9 8.2-5.5c2.1-1.4 4.2-2.7 6.3-3.9-9.2-114.1-5.7-156.3 36.6-204.5 82.6-67.7 272-65.3 355.7-3.9 57 53.8 61 114.1 43.5 207.3zM551.3 644.1l1.2 26.1-15.5 25.6 21.6 141.8L647 647.5l134.7-4.6c69.6 65.8 114.2 220.8 103.2 321.9H134.2c1.9-88.8 18.2-240.4 106.6-317.6l121.7 1.1 114 188 21.4-140.6-15.5-25.5 1.2-26.1c29.5-1.6 38.2-1.6 67.7 0zM652 273.7c-53.5 10.5-133.2 19.6-196.2-15.8-24.2-13.6-59.7 14.3-88.7 11.4-9 17.8-15.7 37.3-19.6 58l-3.2 17.1-17.2-1.7c-3.2-0.3-6.5 0.1-10 1.3-1.6 0.5-3.1 1.1-4.7 1.9 0.5 19 3.1 34.7 8 46.5 4.3 10.4 10.4 17.1 18.2 19.6l10.1 3.1 2.8 10.1c8.4 31.3 17.9 60.8 31.2 84.8 12.3 22.3 28.1 40 50 50.1 9.2 4.2 38.4 6.6 69 7.1 32.5 0.5 64.9-1.3 75.6-5.2 24.1-8.9 41.4-26.1 54.9-48.6 14.5-24.2 24.8-54.9 33.8-88.6l2.6-9.6 9.6-3.3c7.6-2.7 13.5-9.6 17.6-20 4.7-11.7 7.3-27.2 7.7-45.8-1.4-0.7-2.8-1.3-4.2-1.7-3.4-1.1-6.6-1.7-9.7-1.5l-16.8 1.1-3.1-16.5c-3.8-19.2-9.7-37.2-17.7-53.8z m0 0'
female = 'path://M354.304 631.808c18.944-7.168 44.544-17.408 56.32-26.624 7.168-5.632 11.264-16.384 12.8-28.16-4.608-3.072-9.216-6.144-12.288-9.728-53.248-54.784-56.32-132.608-56.832-229.376 8.704 1.536 271.872 5.12 284.16 3.584 2.048 109.056-5.632 176.128-68.608 234.496l-1.024 1.024c2.048 12.288 6.144 23.04 12.8 28.16 12.8 9.728 35.84 20.992 51.2 28.16-29.696 108.544-66.048 172.544-103.936 195.584-5.632-20.48-6.144-44.544-0.512-72.192-4.608-2.048-9.216-3.584-14.336-5.632-5.12 6.656-9.728 13.312-13.824 20.48-4.096-6.656-8.704-13.824-13.824-20.48-4.608 2.048-9.216 3.584-14.336 5.632 5.632 29.184 5.12 54.272-1.536 75.776-41.984-19.968-82.944-88.064-116.224-200.704z m393.216 16.384c-9.216-28.672-16.384-64.512-9.728-95.232 16.384-74.24 3.584-138.752-5.12-159.744-20.992-51.2 2.56-101.888-6.656-140.8-14.848-62.464-41.984-105.984-86.016-135.68-60.928-40.96-151.552-37.888-230.4-20.48-69.12 14.848-134.656 49.152-130.56 107.008 3.072 41.984-4.608 71.168-19.968 92.16-13.824 19.456-15.872 39.424-12.8 60.416 9.216 54.272 48.128 83.968 26.624 136.704-18.944 47.616-41.984 124.928-44.544 169.472-31.232 11.776-66.56 26.112-103.936 39.936-28.672 52.736-38.4 125.952-32.768 268.8H890.88c2.56-111.616 1.024-221.184-40.448-296.96-31.744-7.168-68.608-16.896-102.912-25.6z'
#cube = 'path://M42.666667 128m53.333333 0l832 0q53.333333 0 53.333333 53.333333l0 661.333334q0 53.333333-53.333333 53.333333l-832 0q-53.333333 0-53.333333-53.333333l0-661.333334q0-53.333333 53.333333-53.333333Z'
cube = 'path://M62 62h900v900h-900v-900z'
box = 'path://M962 802H 62V 222h 900v 580z      M162 702h 793V 322H 70v 380z'
boy='path://M516.769444 0c-282.75 0-512 229.25-512 512 0 282.75 229.25 512 512 512 282.75 0 512-229.25 512-512 0-282.875-229.25-512-512-512z m0 987.375c-262.5 0-475.375-212.875-475.375-475.375 0-126.125 50.125-247 139.25-336.125S390.769444 36.5 516.769444 36.5c262.625 0 475.375 212.875 475.375 475.375 0.125 262.625-212.75 475.5-475.375 475.5z M788.144444 691.625c2.5 8 4.25 16.25 5.125 24.625 1.625 10 2.375 20 2.375 30.125 0.125 9.625-0.875 19.25-2.75 28.75-1 7.25-3.625 14.25-7.625 20.375-5.5 4.625-12.125 7.875-19.125 9.375-12.75 3.5-25.625 6.25-38.625 8.25-15.5 2.5-32.625 5-51.75 7.25-19.125 2.25-37.875 4.25-57 5.875-19.125 1.625-37.5 2.75-55.375 3.5-17.875 0.75-33.5 1-47 1s-29.125-0.5-47-1.375c-17.875-0.875-35.875-2.125-54.625-3.5-18.75-1.375-37.5-3-55.75-4.875-18.375-1.875-35.5-4-50.625-6.25s-28.25-4.5-39-6.5c-6.875-0.75-13.625-3-19.5-6.625-6.875-8.75-10.625-19.625-10.75-30.75-2.125-22.5-1-45.25 3.25-67.375 1.625-14.375 9.125-27.375 20.75-35.875 11.75-8.625 24.75-15.25 38.625-19.625 14.75-4.875 30.25-9.125 46.625-12.75 14.875-3.125 29.125-8.75 42.25-16.625 7.875-4.375 15.25-9.75 21.875-16 4.875-4.25 8.75-9.375 11.5-15.25 2-5 3.125-10.5 3.125-16s-0.375-11.75-0.75-18.625c-0.25-9-4.125-17.5-10.75-23.5-6.75-5.875-13.75-11.625-20.75-17.25-4.25-3.125-7.625-7.25-10-12-2.875-5.375-5.5-10.75-8-16.25-2.625-6.5-4.875-13.125-6.75-19.875-0.75-0.875-2.375-2.5-4.375-4.875-1.875-3-3.625-6.125-5.125-9.375-2.625-5.5-5-11.125-6.75-17-2.125-6.5-3.75-13.125-4.75-19.875-0.875-4.875-0.75-10 0.375-14.875 0.25-4.625 1.5-9 3.625-13.125 0.125-16.25 1.125-32.375 3.125-48.5 2.25-14.75 5.5-29.25 9.625-43.5 4.25-15.375 11.5-29.75 21.5-42.125 8.5-10.875 18.375-20.5 29.5-28.75 10-7.5 21-13.375 32.625-17.625 11-4.125 22.25-7 33.875-8.625 11-1.5 22-2.375 33.125-2.375 9.375 0.625 18.75 2 27.875 4 6.125 1.25 12 3.125 17.5 5.875 4.375 2.125 8.375 5 12 8.375 3.125 3.125 6.375 6.5 10 9.625h14.375c4.5 0 8.875 0.625 13.125 1.75 4.625 1.375 8.875 3.625 12.75 6.625 5 4.125 9.5 8.75 13.5 13.875 11.25 13.375 19.875 28.75 25.5 45.25 5.25 15.25 9.25 30.75 12 46.625 2.5 16.75 4.125 33.625 4.75 50.5 1 2.625 1.625 5.5 2 8.25 0.375 3.25 0.75 7.125 1.25 11.75 0.375 5.75-0.125 11.625-1.25 17.25-0.5 7.375-2 14.75-4.375 21.75-1.5 4.375-3.875 8.375-6.75 12-2.125 3-5 5.375-8.375 6.875-1.75 6.875-4 13.5-6.75 19.875-2.375 5.5-4.75 11-7.625 16.25-2.5 4.75-6 8.875-10.375 12-8.375 6.5-15.5 12-20.75 16.25-6.25 6-10.375 13.875-11.5 22.5-1.25 6.375-1.625 12.875-1.25 19.375 0.5 7 2.5 13.75 5.625 19.875 3.75 7.375 9 13.875 15.5 19 9.125 7 19.375 12.375 30.25 16 11.5 4 24.25 8 37.875 11.5 13.5 3.5 27.125 7.375 39.875 11.75 12.125 4.125 23.625 9.875 34.25 17 9.5 6.375 16.75 15.875 20.625 26.625z m-240.375-22.75c1.5-2.625 1.625-5.875 0.375-8.625-1.375-3.125-3.25-6.125-5.625-8.625-2.75-2.75-5.625-5.75-9.125-9h-34.25c-3.125 2.75-5.875 5.75-8.375 9-2 2.25-3.625 4.875-4.75 7.625-1.25 2.375-1.125 5.375 0.375 7.625 3.25 6.5 6.375 12 8.75 17 2 4.375 4.875 8.25 8.375 11.375-1.25 4-2.375 10-4.375 17.625s-4 15.5-6 23.5c-2 8.125-3.125 15.75-4.375 22.875-1 5.375-1.5 10.75-1.625 16.25 0.125 3 1.25 6 3.25 8.25 2.5 3.5 5.125 6.875 8 10 2.625 3.25 5.75 6.125 9.125 8.625 4.875 4.625 12.625 4.625 17.5 0 3.75-2.375 7.125-5 10.375-8 3.125-2.875 5.75-6.125 8-9.625 2-2.5 3.25-5.5 3.625-8.625-0.375-4.75-1.125-9.5-2.375-14.125-1.625-6.75-3.125-14.125-4.75-22.5-1.625-8.25-3.625-16.375-5.625-24.25-2-7.75-3.625-14.125-5.125-18.75 3.25-2.875 6.125-6.25 8.375-10 2.75-4.375 6.25-10.25 10.25-17.625z'
girl='path://M516.769444 0c-282.75 0-512 229.25-512 512 0 282.75 229.25 512 512 512 282.75 0 512-229.25 512-512 0-282.75-229.125-512-512-512z m0 987.375c-262.5 0-475.375-212.75-475.375-475.375C41.394444 385.875 91.519444 265 180.644444 175.875S390.769444 36.625 516.769444 36.625c262.625 0 475.375 212.875 475.375 475.375 0.125 262.625-212.75 475.375-475.375 475.375z M788.019444 687c2.625 7.875 4.5 15.875 5.625 24.125 1.25 9.875 1.875 19.75 2 29.625 0 9.5-0.75 19-2.375 28.375-0.875 7.25-3.5 14.25-7.625 20.375-5.625 4.75-12.25 7.875-19.5 9.25-10.375 3.125-23.875 6.375-39.875 9.5-15.875 3.125-33.875 6.25-53.375 9.25-19.5 3-39.375 5.5-59 7.625-19.5 2.125-38.625 3.75-57 4.875-18.375 1.125-33.875 1.625-47.375 1.625-13.5 0-28.625-0.625-46.25-2-17.5-1.25-34.625-2.875-53.375-4.75-18.75-1.875-36.25-4-54.125-6.75-17.875-2.75-34.25-5.125-49-8-14.75-2.75-27.125-5.25-37-7.625-6.875-1.125-13.5-3.375-19.5-6.75-7-8.75-10.875-19.625-11.125-30.875-2.125-22.375-1-45 3.125-67.125 2.125-14 9.75-26.625 21.125-35 11.875-8.375 25.125-14.875 39-19.125 14.75-4.625 30.25-8.875 46.25-12.75 14.875-3.25 29.125-8.875 42.25-16.5 7.5-4.125 14.625-9 21.125-14.625 4.5-3.875 8.25-8.5 11.125-13.625 2.25-4.25 3.625-8.875 4.375-13.625 0.75-5.125 0.875-10.25 0.375-15.25l-30.625-4.625h-119.5c20.5-22.25 34.375-49.875 39.875-79.625-0.25-59.375 8.375-118.375 25.5-175.25 9.5-21.5 23.375-40.75 40.625-56.625 33.875-36.625 83.625-54.125 133-47 64.125-7.25 126.5 24 159.25 79.625 24.875 63.5 38.25 130.875 39.875 199.125 0.875 31.125 15.5 60.25 39.875 79.625h-119.5l-31.875 4.75c-0.375 5.625 0 11.375 1.25 16.875 1.375 6 3.75 11.75 7.125 16.875 4.5 6.125 9.875 11.5 15.875 15.875 8.5 5.75 17.75 10.25 27.5 13.375 12.125 4.5 24.5 8.375 37 11.5 13.5 3.25 26.75 7.125 39.875 11.75 12.125 4.125 23.625 9.75 34.25 16.875 9.625 6.375 16.875 15.75 20.75 26.625z m-266.5 102.375c3.75-0.5 7.25-1.875 10.375-4 4.875-2.875 9.5-6.25 14-9.875 4.375-3.5 8.5-7.375 12.375-11.375 2.5-2.375 4.375-5.5 5.125-8.875-0.375-4.5-1-9-2-13.375-1.625-6.375-3.125-13.375-5.125-21-2-7.625-3.625-15-5.625-22.25-1.25-6-2.875-12-4.75-17.75 4.625-3.25 8.75-7.125 12.375-11.5 3-3.5 4.875-8 5.125-12.75-0.125-3.75-1.5-7.375-4-10.125-2.625-3.5-5.5-6.625-8.75-9.5-3.625-3-8-5.875-12.375-8.875h-35c-4 3-8 6-11.125 8.875-2.875 2.75-5.5 5.875-7.625 9.25-2.25 3-3.5 6.75-3.625 10.5 0.625 3.625 2.375 6.875 4.75 9.5 3.875 4.625 8.125 8.875 12.75 12.75-2 5.25-3.75 10.75-5.125 16.25-2 7-4.375 14.375-6.75 22.25-2.375 7.875-4.375 15.25-6 22.25-1.375 5.125-2.125 10.375-2.375 15.625 0.875 3.375 2.625 6.5 5.125 8.875 4 4.125 8.25 7.875 12.75 11.375 4.625 3.75 9.5 7 14.75 9.875 3.25 2.125 6.875 3.5 10.75 3.875z'
location='path://M528 32C325.056 32 160 196.8 160 399.36c0 75.2 22.656 147.584 65.024 208.48 2.112 3.648 4.256 7.168 6.784 10.592l268.608 353.472c7.296 8.096 17.088 12.576 27.584 12.576 10.368 0 20.224-4.512 28.768-14.08l267.36-352c2.624-3.52 4.896-7.36 6.112-9.6A364.864 364.864 0 0 0 896 399.36C896 196.8 730.912 32 528 32z m0 498.72a131.52 131.52 0 0 1-131.456-131.232 131.488 131.488 0 0 1 262.88 0 131.52 131.52 0 0 1-131.424 131.2z'
airplane='path://M965.40416 342.76608c-19.50976-52.24704-47.744-100.41472-83.92064-143.16288-18.26688-21.58592-50.57408-24.27392-72.15872-6.0096-21.58464 18.26688-24.2752 50.5728-6.00832 72.15872C861.4336 334.42816 893.44 421.88032 893.44 512c0 210.32576-171.11424 381.44-381.44 381.44-210.32704 0-381.44-171.11424-381.44-381.44 0-210.32704 171.11296-381.44 381.44-381.44 60.48896 0 118.30784 13.72672 171.85152 40.8 25.2288 12.75648 56.03328 2.64704 68.7936-22.58816 12.76032-25.23392 2.64576-56.03456-22.58816-68.7936C662.05312 45.5936 588.68992 28.16 512 28.16c-65.29536 0-128.66432 12.8-188.34688 38.04416-57.62176 24.37248-109.36064 59.2512-153.7792 103.66976-44.41856 44.41856-79.29856 96.15744-103.66976 153.77792C40.96 383.3344 28.16 446.70464 28.16 512c0 65.29536 12.8 128.6656 38.04416 188.34688 24.37248 57.62176 59.2512 109.36064 103.66976 153.7792s96.15744 79.29856 153.7792 103.66976C383.33568 983.04 446.70464 995.84 512 995.84c65.29536 0 128.6656-12.8 188.34816-38.04416 57.62176-24.3712 109.36064-59.2512 153.77792-103.66976 44.41856-44.41856 79.29856-96.15744 103.66976-153.7792C983.04 640.6656 995.84 577.29536 995.84 512 995.84 453.79328 985.6 396.85504 965.40416 342.76608z  M718.52672 248.7872c-28.15104 3.3984-50.06336 18.60608-68.33152 39.93472-17.31456 20.21376-35.4048 39.76192-53.09312 59.65568-9.1136 10.24896-17.9328 20.76672-27.22688 30.848-3.22944 3.50464-5.80992 8.89984-12.31232 7.28448-3.84896-0.95616-7.7504-1.69984-11.59936-2.65472-36.48128-9.04064-72.96128-18.08512-109.43232-27.168-60.3712-15.03616-120.73344-30.10816-181.10592-45.14432-9.62304-2.39616-18.24896 0.9856-22.54336 9.56544-3.24864 6.49216-4.54016 14.9376 4.33152 22.69952 5.83808 5.10848 12.08832 9.75232 18.2336 14.5024 19.91424 15.39456 39.84512 30.76864 59.808 46.10048 16.80256 12.90496 33.6896 25.69984 50.48192 38.61632 19.79008 15.22432 39.51104 30.53696 59.2768 45.79328 8.96128 6.91712 17.95584 13.78944 27.40352 21.04064-19.81568 21.91744-39.24864 43.34848-58.5984 64.85376-4.78592 5.31968-8.928 11.26912-14.04416 16.22784-5.04192 4.88704-10.34752 11.35488-16.57088 12.86272-32.064 7.76576-64.45824 14.1632-96.74112 21.02912-8.75776 1.8624-13.8624 8.44928-14.24256 17.36192-0.544 12.77568 5.64864 16.6592 15.6544 20.66688 25.02016 10.0224 50.08896 19.92704 75.02592 30.15552 2.44352 1.00224 4.96896 3.5456 5.9968 5.98784 11.0208 26.20032 21.80864 52.50048 32.52992 78.82496 4.96 12.18176 18.5408 14.64064 27.68256 9.10848 7.41632-4.48896 8.10752-11.79648 9.67168-19.33568 5.04576-24.3264 10.70976-48.52608 16.34048-72.72704 2.4192-10.39616 2.5728-21.76256 12.8256-28.92416 5.67168-3.9616 10.81984-8.71168 15.99488-13.34016 21.27232-19.02336 42.45248-38.15168 63.70816-57.19424 1.73824-1.55776 3.80032-2.752 5.95072-4.288 2.752 3.44064 5.02272 6.20288 7.21152 9.02784 14.5216 18.73664 29.00352 37.50272 43.5392 56.22784 10.68288 13.76256 21.39776 27.49952 32.15232 41.20448 14.64832 18.66496 29.36576 37.27488 44.02304 55.93344 14.13248 17.9904 28.14464 36.07424 42.3808 53.97888 4.43392 5.57696 14.08384 8.33408 20.02688 5.9008 6.4384-2.6368 14.7328-8.84864 12.9088-18.06464-1.53728-7.76576-3.0528-15.55328-5.02528-23.21664-11.86048-46.09408-23.89504-92.14208-35.7824-138.22976-12.672-49.12256-25.27744-98.26176-37.76256-147.43168-0.48256-1.90208-0.07552-4.68096 1.056-6.208 2.2592-3.04768 5.19168-5.64224 8.05376-8.18944 18.61248-16.56704 37.248-33.1072 55.94496-49.57952 9.17248-8.08064 18.71104-15.76064 27.68384-24.05376 8.83712-8.16768 17.97888-16.22144 25.55648-25.49376 16.64768-20.36992 22.62272-44.4032 20.3456-70.34624-1.32224-15.0592-7.29728-28.56576-18.59968-38.81472C760.5312 248.56064 739.60704 246.24256 718.52672 248.7872zM761.84064 323.68128c1.86496 19.5456-6.144 33.90208-21.51552 44.9664-3.6032 2.59456-7.03616 2.65856-10.54464-0.34688-8.30976-7.11808-16.67712-14.17088-25.08672-21.17376-9.52448-7.93344-19.16416-15.72864-28.6592-23.69792-3.52768-2.96064-6.2976-5.32608-3.69536-11.50208 6.65472-15.78752 25.36192-30.99392 43.86688-29.17632C737.91744 281.6704 759.61984 300.41984 761.84064 323.68128z'
star = 'path://M956.741665 419.885046c-5.241374-16.200984-19.178805-28.054963-36.034704-30.496571l-254.931157-44.791136L551.77233 115.15981c-7.444553-15.248286-22.992667-24.898063-40.025599-24.898063-17.036002 0-32.522718 9.648754-40.146349 24.898063L357.655238 344.597339 102.784456 390.936737c-16.856923 2.442632-30.852682 14.295588-36.035728 30.496571-5.300726 16.141632-0.893346 33.951253 11.317767 45.804209l184.407963 170.410158-48.186466 247.662613c-2.859118 16.737196 3.930519 33.833573 17.809621 43.720757 13.244652 13.244652 40.562834 6.90834 47.053666 3.514033l232.595452-113.64634L742.85225 932.545078c6.490831 3.394307 32.166607 11.373025 47.052642-3.514033 13.760398-9.887184 20.787442-26.982538 17.869996-43.720757L761.077334 637.647675l184.467315-172.018795C957.69641 453.836299 962.043415 436.026678 956.741665 419.885046z'